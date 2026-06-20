#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Erzeugt den Prüfbericht zu Aufgabe 4 als professionelles Word-Dokument:
  - Screenshots der gefundenen Fehler (bugs/BUG-01.png, bugs/BUG-02.png)
  - Aufgabe4_Pruefbericht.docx  (Deckblatt, Inhaltsverzeichnis, nummerierte Kapitel,
    formatierte/farbcodierte Tabellen mit allen Pflichtspalten, eingebettete Screenshots,
    Kopf-/Fußzeile mit Seitenzahl)

Tabellen-Daten kommen aus testkonzept-aufgabe4.md + testfaelle-aufgabe4.md (Single Source).

Setup:
    python3 -m venv /tmp/a4venv && /tmp/a4venv/bin/pip install Pillow python-docx
    /tmp/a4venv/bin/python build-docx.py
"""
import os
import re
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
FONT = "/System/Library/Fonts/Menlo.ttc"
USABLE_MM = 186.0  # A4 Hochformat, 12 mm Ränder

# ── Stil (aus Design-Spezifikation) ──────────────────────────────────────────
C_HEAD = RGBColor(0x1F, 0x4E, 0x79)
HEX_HEADER_FILL = "1F4E79"
HEX_HEADER_TEXT = "FFFFFF"
HEX_ZEBRA = "F2F6FB"
HEX_OK = "1E7E34"
HEX_FAIL = "C0392B"
HEX_BLOCK = "B9770E"
HEX_FAILROW = "FBEAEA"
TABLE_PT = 7.5
TOK = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]*\)|\*[^*]+\*)")


# ── Screenshots ──────────────────────────────────────────────────────────────
def render_terminal_png(txt_path, out_path, title):
    with open(txt_path, encoding="utf-8") as fh:
        lines = [l.rstrip("\n") for l in fh if l.strip() != ""] or [" "]
    FS, PAD, TB = 22, 26, 44
    LH = int(FS * 1.5)
    font = ImageFont.truetype(FONT, FS)
    m = ImageDraw.Draw(Image.new("RGB", (4, 4)))
    W = int(max(m.textlength(l, font=font) for l in lines)) + PAD * 2
    H = TB + PAD + LH * len(lines) + PAD
    img = Image.new("RGB", (W, H), "#1e1f29")
    d = ImageDraw.Draw(img)
    d.rectangle([0, 0, W, TB], fill="#2d2f3a")
    for i, c in enumerate(["#ff5f56", "#ffbd2e", "#27c93f"]):
        x = PAD + i * 28
        d.ellipse([x, TB / 2 - 7, x + 14, TB / 2 + 7], fill=c)
    d.text((W / 2, TB / 2), title, fill="#c8c8d0", font=font, anchor="mm")
    y = TB + PAD
    for l in lines:
        col = "#d6d6dd"
        if l.lstrip().startswith("#"):
            col = "#f1c40f"
        elif l.startswith(">>>"):
            col = "#5bc0eb"
        elif "FEHLER" in l:
            col = "#ff6b6b"
        elif "HTTP-Status:" in l:
            col = "#ffa657"
        d.text((PAD, y), l, fill=col, font=font)
        y += LH
    img.save(out_path)
    print("  PNG:", os.path.relpath(out_path, HERE), "(%dx%d)" % (W, H))


# ── Markdown-Tabellen einlesen ───────────────────────────────────────────────
def parse_all_tables(text):
    lines = text.split("\n")
    out, heading, i, n = [], "", 0, len(lines)
    while i < n:
        s = lines[i].strip()
        if s.startswith("#"):
            heading = s.lstrip("#").strip()
            i += 1
            continue
        if s.startswith("|"):
            block = []
            while i < n and lines[i].strip().startswith("|"):
                block.append(lines[i].strip())
                i += 1
            rows = [[c.strip() for c in b.strip("|").split("|")] for b in block]
            sep = 2 if len(rows) > 1 and re.fullmatch(r"[\s:\-|]+", "|".join(rows[1])) else 1
            out.append({"heading": heading, "cols": rows[0], "rows": rows[sep:]})
            continue
        i += 1
    return out


def get_table(tables, *subs, n=0):
    hits = [t for t in tables if all(x.lower() in t["heading"].lower() for x in subs)]
    return hits[n]


# ── Inline-Formatierung ──────────────────────────────────────────────────────
def add_runs(p, text, size=None, color=None):
    for part in TOK.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1]); r.font.name = "Consolas"
        elif part.startswith("[") and "](" in part:
            r = p.add_run(re.match(r"\[([^\]]+)\]", part).group(1))
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            r = p.add_run(part[1:-1]); r.italic = True
        else:
            r = p.add_run(part)
        if size:
            r.font.size = size
        if color:
            r.font.color.rgb = color


def clean_inline(t):
    t = re.sub(r"\*\*([^*]+)\*\*", r"\1", t)
    t = re.sub(r"`([^`]+)`", r"\1", t)
    t = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", t)
    return t.strip()


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement("w:tblHeader")
    h.set(qn("w:val"), "true")
    trPr.append(h)


def set_widths(table, widths_mm):
    table.autofit = False
    table.allow_autofit = False
    tblPr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    grid = table._tbl.tblGrid
    for i, w in enumerate(widths_mm):
        dxa = str(int(w * 1440 / 25.4))
        grid[i].set(qn("w:w"), dxa)
        for row in table.rows:
            row.cells[i].width = Mm(w)


def distribute(cols, total=USABLE_MM):
    weights = [max(12, len(c)) for c in cols]
    s = sum(weights)
    return [round(total * w / s, 1) for w in weights]


def styled_table(doc, cols, rows, widths=None, font_pt=TABLE_PT, result_idx=None):
    t = doc.add_table(rows=1, cols=len(cols))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for j, c in enumerate(cols):
        hdr[j].text = ""
        add_runs(hdr[j].paragraphs[0], clean_inline(c), size=Pt(font_pt),
                 color=RGBColor(0xFF, 0xFF, 0xFF))
        for r in hdr[j].paragraphs[0].runs:
            r.bold = True
        shade(hdr[j], HEX_HEADER_FILL)
    repeat_header(t.rows[0])
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        is_fail = result_idx is not None and "fehlerhaft" in clean_inline(
            row[result_idx] if result_idx < len(row) else "").lower()
        for j in range(len(cols)):
            val = row[j] if j < len(row) else ""
            cell = cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            if result_idx is not None and j == result_idx:
                txt = clean_inline(val)
                add_runs(p, txt, size=Pt(font_pt), color=RGBColor(0xFF, 0xFF, 0xFF))
                for r in p.runs:
                    r.bold = True
                low = txt.lower()
                shade(cell, HEX_FAIL if "fehler" in low else HEX_BLOCK if "block" in low else HEX_OK)
            else:
                add_runs(p, val, size=Pt(font_pt))
                if is_fail:
                    shade(cell, HEX_FAILROW)
                elif ri % 2 == 1:
                    shade(cell, HEX_ZEBRA)
    set_widths(t, widths or distribute(cols))
    doc.add_paragraph()
    return t


# ── Felder / Kopf-/Fußzeile / TOC ────────────────────────────────────────────
def add_field(paragraph, instr, placeholder="1"):
    run = paragraph.add_run()
    for typ in ("begin",):
        e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), typ); run._r.append(e)
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    run._r.append(it)
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate"); run._r.append(sep)
    t = OxmlElement("w:t"); t.text = placeholder; run._r.append(t)
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end"); run._r.append(end)


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin"); run._r.append(b)
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
    it.text = 'TOC \\o "1-2" \\h \\z \\u'; run._r.append(it)
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate"); run._r.append(sep)
    t = OxmlElement("w:t")
    t.text = "Inhaltsverzeichnis: in Word mit „Alles markieren“ (Cmd+A) und F9 aktualisieren."
    run._r.append(t)
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end"); run._r.append(end)


def enable_update_fields(doc):
    s = doc.settings.element
    el = OxmlElement("w:updateFields"); el.set(qn("w:val"), "true"); s.append(el)


# ── Inhalte (aus Design-Spezifikation, self-contained) ───────────────────────
SECTIONS = [
 ("1", "Einleitung & Ziel des Berichts",
  "Dieser Prüfbericht weist den spezifikationsbasierten (Black-Box-)Testfallentwurf, dessen reale "
  "Durchführung per Postman/Newman sowie den Nachweis der gefundenen Fehler für genau einen "
  "Anwendungsfall nach. Als Einzelarbeit wird gemäß Vorgabe ein repräsentativer Anwendungsfall "
  "behandelt. Belegt werden (1.1) die Ableitung der Testfälle aus fünf spezifikationsbasierten "
  "Entwurfsmethoden, (1.2) die dokumentierte Ausführung aller Testfälle und (1.3) der Fehlernachweis "
  "mit eingebetteten Screenshots. Abgrenzung zu Aufgabe 3: Dort wurde die Überstunden-Berechnung per "
  "White-Box-Entscheidungsüberdeckung geprüft; Aufgabe 4 prüft komplementär einen anderen Use Case "
  "rein spezifikationsbasiert."),
 ("2", "Pflichtspalten-Nachweismatrix",
  "Die folgende Matrix bildet jedes der 15 von Aufgabe 4 geforderten Pflicht-Attribute auf seine "
  "exakte Fundstelle im Bericht ab. Die Entwurfs-Attribute (1–8) liegen in Tabelle 1, die "
  "Ausführungs-Attribute (9–15) in Tabelle 2; beide sind über die eindeutige Testfall-ID verknüpft. "
  "Die Screenshots der Fehler sind in Anhang A eingebettet."),
 ("3", "Anwendungsfall & Auswahlbegründung",
  "Geprüft wird der Anwendungsfall „Zeiterfassung starten / beenden“: Der Arbeiter stempelt ein "
  "(POST /zeit/start → Zustand LAUFEND), aus (POST /zeit/stop → BEENDET) und korrigiert eine Buchung "
  "bei Bedarf (PATCH /zeit/{id}). Die Auswahl begründet sich durch das höchste Produktrisiko (in "
  "Aufgabe 2 ist „Falsche Erfassung von Arbeitszeiten“ mit Schaden hoch bewertet und fließt direkt in "
  "Lohnabrechnung und Finanzamtsbelege ein), die reichste spezifikationsbasierte Angriffsfläche "
  "(echter Zustandsautomat, zahlreiche Grenzwerte, Entscheidungslogik für die Pausenberechnung) sowie "
  "die bewusste Abgrenzung zu Aufgabe 3. Da die Anforderungen in Aufgabe 1 als Anwendungsfälle und "
  "nicht als User Stories beschrieben wurden, erfolgt der Entwurf anwendungsfallbasiert."),
 ("4", "Testbasis (Spezifikation)",
  "Quelle der Testbasis sind der REST-API-Vertrag und die Geschäftsregeln des Backends. Die folgenden "
  "Endpunkte des Anwendungsfalls (alle unter /api/v1, JWT-Bearer) bilden die Schnittstelle; darunter "
  "stehen die neun als Testbasis verwendeten Geschäftsregeln."),
 ("5", "Spezifikationsbasierte Entwurfsmethoden",
  "Für jede der fünf Black-Box-Methoden werden Prinzip und konkrete Anwendung auf diesen Anwendungsfall "
  "gezeigt, jeweils mit Verweis auf die resultierenden Testfall-IDs."),
 ("6", "Rahmendaten der Durchführung & Rückverfolgbarkeit",
  "Die folgenden Ausführungs-Metadaten gelten konstant für alle 31 Testfälle. Anschließend bildet die "
  "Rückverfolgbarkeitsmatrix die Testziele aus Aufgabe 2 auf die abdeckenden Testfälle ab."),
 ("7", "Tabelle 1 — Testfallentwurf",
  "Vollständige Entwurfstabelle aller 31 Testfälle mit den Pflicht-Entwurfsspalten ID, Testfall-Titel, "
  "Autor, Getestete Anforderung, Vorbedingung, Benötigte Testdaten, Grober Testablauf und Erwartetes "
  "Ergebnis, ergänzt um die Spalte Methode. Methoden-Legende: ÄK = Äquivalenzklassen, GW = "
  "Grenzwertanalyse, ET = Entscheidungstabelle, ZÜ = Zustandsübergang, AF = Anwendungsfalltest."),
 ("8", "Tabelle 2 — Testausführung (Erweiterung, verknüpft über ID)",
  "Erweiterung derselben 31 Testfälle um die Pflicht-Ausführungsspalten Datum, Tester, Version der "
  "Software, Umgebung, Ergebnis (farbcodiert), Gefundene Abweichung(en) und Kommentare. Die beiden "
  "fehlerhaften Testfälle TC-STOP-07 und TC-PATCH-09 sind rot hervorgehoben und verweisen auf die "
  "Screenshots in Anhang A."),
 ("9", "Ergebnisübersicht & gefundene Fehler (Defect-Register)",
  "Kennzahlenbilanz: 31 Testfälle gesamt · 29 erfolgreich · 2 fehlerhaft · 0 blockiert (Newman: 54 "
  "Requests inkl. Setup, 88 Assertions, 2 Fehlschläge). Das Defect-Register listet die gefundenen "
  "Fehler mit Schwere und Nachweis. Korrekturempfehlung: die Zukunfts-Validierung analog zu start_zeit "
  "auch für end_zeit erzwingen (BUG-01) und pause_minuten auf laufenden Buchungen ablehnen (BUG-02)."),
 ("10", "Anhang A — Screenshots der gefundenen Fehler",
  "Anhang A bettet die zwei Konsolen-Screenshots der real reproduzierten Fehler ein. Jeder Screenshot "
  "zeigt Request, HTTP-Status und Antwort im Schema „Erwartet 400 vs. Ist 200“. Die Abbildungsnummern "
  "sind wortgleich mit den Referenzen in Tabelle 2 (Spalte „Gefundene Abweichung(en)“) und im "
  "Defect-Register."),
 ("11", "Anhang B — Reproduktion & Artefakte",
  "Alle Artefakte liegen reproduzierbar im Projekt. Die Testfälle lassen sich mit den genannten "
  "Befehlen erneut ausführen; das Ergebnis ist deterministisch (29 erfolgreich, 2 fehlerhaft)."),
]

GESCHAEFTSREGELN = [
 "Pro Arbeiter darf nur EINE laufende Buchung existieren (sonst HTTP 409).",
 "end_zeit muss echt nach start_zeit liegen.",
 "pause_minuten ≥ 0 und ≤ Spanne (end − start in Minuten).",
 "start_zeit darf nicht in der Zukunft liegen.",
 "baustelle_id ist optional, muss aber existieren, wenn angegeben (Fremdschlüssel).",
 "notiz ≤ 1000 Zeichen.",
 "Zeiten werden in UTC gespeichert; Zeitzonen-Offsets werden normalisiert.",
 "Gesetzliche Pause (§11 AZG): Spanne > 360 min ⇒ 30 min Pause; dauer = Spanne − pause (≥ 0).",
 "Mandantentrennung: Arbeiter sehen/ändern nur eigene Buchungen; /admin/* nur für Rolle admin.",
]

STATE_DIAGRAM = [
 "          start (201)                stop (200)",
 "KEINE  ───────────────▶  LAUFEND  ───────────────▶  BEENDET",
 "  ▲                        │   ▲                        │",
 "  │  stop (409, verboten)  │   │ patch (200, LAUFEND)   │ patch (200, BEENDET)",
 "  │                        │   ╰────────────────────────╯",
 "  ╰── start auf LAUFEND → 409 ; stop auf KEINE/BEENDET → 409 (verboten)",
]

ARTIFACTS = [
 "Postman-Collection:  postman/zeiterfassung.postman_collection.json",
 "Postman-Environment: postman/local.postman_environment.json",
 "Durchführungsbericht: durchfuehrungsbericht/ (Newman JSON, HTML, Text)",
 "Repro-Skripte:        bugs/BUG-01-future-end-zeit.sh , bugs/BUG-02-pause-auf-laufender-buchung.sh",
 "Screenshots:          bugs/BUG-01.png , bugs/BUG-02.png",
 "Testkonzept/Testfälle (Quelle): testkonzept-aufgabe4.md , testfaelle-aufgabe4.md",
 "Reproduktion: make db-up && make migrate-up && make seed && make run-api ; danach "
 "newman run postman/zeiterfassung.postman_collection.json -e postman/local.postman_environment.json",
]

CAP1 = ("Abbildung 1 — BUG-01: end_zeit in der Zukunft (2030) wird bei Stop UND PATCH akzeptiert "
        "(HTTP 200 statt 400); berechnete Dauer ≈ 1.858.410 min (≈ 1290 Tage). Schwere mittel–hoch — "
        "betrifft das Top-Produktrisiko Arbeitszeit/Lohn. Testfall TC-STOP-07. Quelle: bugs/BUG-01.png")
CAP2 = ("Abbildung 2 — BUG-02: pause_minuten auf einer laufenden Buchung wird akzeptiert (HTTP 200 statt "
        "400) und beim Stop still verworfen. Schwere niedrig. Testfall TC-PATCH-09. Quelle: bugs/BUG-02.png")

COVERAGE = [
 ("(1) Eindeutige ID", "Tabelle 1 & 2, Spalte „ID“ (TC-START-01 … TC-AUTH-03)"),
 ("(2) Testfall-Titel", "Tabelle 1, Spalte „Testfall-Titel“"),
 ("(3) Autor", "Tabelle 1, Spalte „Autor“ (L. Prekaj)"),
 ("(4) Getestete Anforderung (Anwendungsfall/Story)", "Tabelle 1, Spalte „Getestete Anforderung“"),
 ("(5) Vorbedingung", "Tabelle 1, Spalte „Vorbedingung“"),
 ("(6) Benötigte Testdaten", "Tabelle 1, Spalte „Benötigte Testdaten“"),
 ("(7) Grober Testablauf", "Tabelle 1, Spalte „Grober Testablauf“"),
 ("(8) Erwartetes Ergebnis", "Tabelle 1, Spalte „Erwartetes Ergebnis“"),
 ("(9) Datum der Durchführung", "Tabelle 2, Spalte „Datum“ + Kapitel 6"),
 ("(10) Tester", "Tabelle 2, Spalte „Tester“ + Kapitel 6"),
 ("(11) Version der Software", "Tabelle 2, Spalte „Version“ + Kapitel 6"),
 ("(12) Umgebung", "Tabelle 2, Spalte „Umgebung“ + Kapitel 6"),
 ("(13) Ergebnis (Erfolgreich/fehlerhaft/blockiert)", "Tabelle 2, Spalte „Ergebnis“ (farbcodiert)"),
 ("(14) Gefundene Abweichung(en)", "Tabelle 2, Spalte „Gefundene Abweichung(en)“ + Kapitel 9"),
 ("(15) Kommentare", "Tabelle 2, Spalte „Kommentare“"),
 ("Screenshots der Fehler", "Anhang A (Abb. 1 = BUG-01, Abb. 2 = BUG-02), referenziert in Tab. 2 & Kap. 9"),
]

# Hochformat: nutzbare Breite 186 mm
T1_WIDTHS = [15, 26, 12, 12, 24, 24, 24, 17, 32]   # ID,Titel,Methode,Autor,Anforderung,Vorbed.,Testdaten,Ablauf,Erwartet — Summe 186
T2_WIDTHS = [15, 16, 13, 15, 26, 18, 41, 42]       # ID,Datum,Tester,Version,Umgebung,Ergebnis,Abweichung,Kommentar — Summe 186


# ── Aufbau ───────────────────────────────────────────────────────────────────
def heading(doc, text, level=1):
    h = doc.add_heading(text, level)
    for r in h.runs:
        r.font.color.rgb = C_HEAD
    return h


def para(doc, text, size=10, italic=False, align=None):
    p = doc.add_paragraph()
    add_runs(p, text, size=Pt(size))
    if italic:
        for r in p.runs:
            r.italic = True
    if align is not None:
        p.alignment = align
    return p


print("Screenshots rendern:")
render_terminal_png(os.path.join(HERE, "bugs", "BUG-01.txt"),
                    os.path.join(HERE, "bugs", "BUG-01.png"),
                    "Terminal  —  BUG-01  (end_zeit in der Zukunft)")
render_terminal_png(os.path.join(HERE, "bugs", "BUG-02.txt"),
                    os.path.join(HERE, "bugs", "BUG-02.png"),
                    "Terminal  —  BUG-02  (Pause auf laufender Buchung)")

tk = parse_all_tables(open(os.path.join(HERE, "testkonzept-aufgabe4.md"), encoding="utf-8").read())
tf = parse_all_tables(open(os.path.join(HERE, "testfaelle-aufgabe4.md"), encoding="utf-8").read())

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10)
sec = doc.sections[0]
sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width, sec.page_height = Mm(297), Mm(210)
sec.left_margin = sec.right_margin = Mm(12)
sec.top_margin = sec.bottom_margin = Mm(14)
sec.different_first_page_header_footer = True

# Kopfzeile (ab Seite 2)
hp = sec.header.paragraphs[0]
add_runs(hp, "Software Testing — Aufgabe 4: Zeiterfassung starten/beenden  ·  Lindrit Prekaj · 2026-06-20",
         size=Pt(8))
for r in hp.runs:
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
# Fußzeile mit Seitenzahl
fp = sec.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_runs(fp, "Version 0cae33d · Lokal (Docker) · Seite ", size=Pt(8))
add_field(fp, "PAGE")
add_runs(fp, " von ", size=Pt(8))
add_field(fp, "NUMPAGES")
for r in fp.runs:
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    r.font.size = Pt(8)

# ── Deckblatt ────────────────────────────────────────────────────────────────
para(doc, "Hochschulmodul Software Testing · Aufgabe 4 · Einzelarbeit", size=9, italic=True,
     align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(2):
    doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Prüfbericht: Spezifikationsbasierter Testfallentwurf & Durchführung")
r.bold = True; r.font.size = Pt(26); r.font.color.rgb = C_HEAD
st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st.add_run("Anwendungsfall: Zeiterfassung starten / beenden — Prekaj-Zeiterfassung (Go-REST-API)")
r.font.size = Pt(15)
hr = doc.add_paragraph()
pPr = hr._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
for k, v in (("w:val", "single"), ("w:sz", "12"), ("w:space", "1"), ("w:color", "1F4E79")):
    bottom.set(qn(k), v)
pbdr.append(bottom); pPr.append(pbdr)
doc.add_paragraph()

meta = [
 ("Autor / Tester", "Lindrit Prekaj"),
 ("Arbeitsform", "Einzelarbeit (1 Anwendungsfall)"),
 ("Datum der Durchführung", "2026-06-20"),
 ("Software-Version", "git 0cae33d (Branch main)"),
 ("Umgebung", "Lokal — macOS (M3) · PostgreSQL 17 (Docker) · Go-API :8080"),
 ("Ausführung / Werkzeug", "Postman-Collection + Newman 6.2.2"),
 ("Testobjekt", "REST-API /api/v1"),
]
mt = doc.add_table(rows=0, cols=2)
mt.alignment = WD_TABLE_ALIGNMENT.CENTER
for k, v in meta:
    cells = mt.add_row().cells
    cells[0].text = ""; rr = cells[0].paragraphs[0].add_run(k); rr.bold = True; rr.font.size = Pt(11)
    cells[1].text = ""; cells[1].paragraphs[0].add_run(v).font.size = Pt(11)
set_widths(mt, [70, 170])
doc.add_paragraph()

badge = doc.add_table(rows=1, cols=4); badge.alignment = WD_TABLE_ALIGNMENT.CENTER
bvals = [("31 Testfälle", "1F4E79"), ("29 Erfolgreich", HEX_OK), ("2 fehlerhaft", HEX_FAIL), ("0 blockiert", HEX_BLOCK)]
for j, (txt, hexc) in enumerate(bvals):
    c = badge.rows[0].cells[j]; c.text = ""
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn = p.add_run(txt); rn.bold = True; rn.font.size = Pt(12); rn.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shade(c, hexc)
set_widths(badge, [60, 60, 60, 60])
para(doc, "Gefundene Fehler: BUG-01 (end_zeit in der Zukunft) · BUG-02 (Pause auf laufender Buchung)",
     size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ── Inhaltsverzeichnis ───────────────────────────────────────────────────────
heading(doc, "Inhaltsverzeichnis", 1)
add_toc(doc)
para(doc, "In Word: Alles markieren (Cmd+A) und F9 drücken, um Verzeichnis und Seitenzahlen zu füllen.",
     size=8, italic=True)
for num, title, _ in SECTIONS:
    para(doc, f"{num}.  {title}", size=10)
doc.add_page_break()

# ── Kapitel ──────────────────────────────────────────────────────────────────
for num, title, body in SECTIONS:
    if num in ("7", "8", "10"):
        doc.add_page_break()
    heading(doc, f"{num}. {title}", 1)
    para(doc, body, size=10)

    if num == "2":
        styled_table(doc, ["Pflicht-Attribut (Aufgabe 4)", "Fundstelle im Bericht"],
                     [[a, o] for a, o in COVERAGE], widths=[95, 178], font_pt=9)
    elif num == "4":
        et = get_table(tk, "Testbasis")
        styled_table(doc, et["cols"], et["rows"], widths=[26, 70, 177], font_pt=9)
        heading(doc, "Geschäftsregeln (Testbasis)", 2)
        for i, reg in enumerate(GESCHAEFTSREGELN, 1):
            add_runs(doc.add_paragraph(style="List Number"), reg, size=Pt(10))
    elif num == "5":
        for sub, label in (("3.1", "5.1 Äquivalenzklassenbildung (ÄK)"),
                           ("3.2", "5.2 Grenzwertanalyse (GW)")):
            heading(doc, label, 2)
            tb = get_table(tk, sub)
            styled_table(doc, tb["cols"], tb["rows"], font_pt=8.5)
        heading(doc, "5.3 Entscheidungstabelle (ET)", 2)
        for k in (0, 1):
            tb = get_table(tk, "3.3", n=k)
            styled_table(doc, tb["cols"], tb["rows"], font_pt=8.5)
        heading(doc, "5.4 Zustandsübergangstest (ZÜ)", 2)
        cp = doc.add_paragraph()
        for k, l in enumerate(STATE_DIAGRAM):
            rr = cp.add_run(l); rr.font.name = "Consolas"; rr.font.size = Pt(8.5)
            if k < len(STATE_DIAGRAM) - 1:
                rr.add_break()
        tb = get_table(tk, "3.4")
        styled_table(doc, tb["cols"], tb["rows"], font_pt=8.5)
        heading(doc, "5.5 Anwendungsfalltest (AF)", 2)
        para(doc, "Hauptpfad (einstempeln → arbeiten → ausstempeln): TC-START-01 → TC-STOP-01 → "
                  "TC-LIST-03. Alternativpfade: ohne Baustelle (TC-START-03), nachträgliche Korrektur "
                  "(TC-PATCH-01/-02). Ausnahmepfade: doppeltes Einstempeln (TC-START-02), Ausstempeln "
                  "ohne laufende Buchung (TC-STOP-02), ungültige Eingaben (TC-START-05, TC-STOP-03/-04), "
                  "fehlende Rechte (TC-AUTH-01/-02).", size=10)
    elif num == "6":
        rd = get_table(tf, "Rahmendaten")
        styled_table(doc, rd["cols"], rd["rows"], widths=[60, 213], font_pt=9)
        heading(doc, "Rückverfolgbarkeit zu den Testzielen (Aufgabe 2)", 2)
        rv = get_table(tk, "Rückverfolgbarkeit")
        styled_table(doc, rv["cols"], rv["rows"], widths=[120, 153], font_pt=9)
    elif num == "7":
        t1 = get_table(tf, "Tabelle 1")
        cols = t1["cols"][:]
        cols.insert(3, "Autor")
        rows = [r[:3] + ["L. Prekaj"] + r[3:] for r in t1["rows"]]
        styled_table(doc, cols, rows, widths=T1_WIDTHS, font_pt=7.5)
    elif num == "8":
        t2 = get_table(tf, "Tabelle 2")
        rows = []
        for r in t2["rows"]:
            r = r[:]
            tcid = clean_inline(r[0])
            if tcid == "TC-STOP-07":
                r[6] = "**BUG-01** → Abb. 1 (Kap. 9, Anhang A)"
            elif tcid == "TC-PATCH-09":
                r[6] = "**BUG-02** → Abb. 2 (Kap. 9, Anhang A)"
            rows.append(r)
        styled_table(doc, t2["cols"], rows, widths=T2_WIDTHS, font_pt=7.5, result_idx=5)
    elif num == "9":
        df = get_table(tf, "Gefundene Fehler")
        styled_table(doc, df["cols"], df["rows"], font_pt=9)
    elif num == "10":
        para(doc, CAP1, size=10)
        doc.add_picture(os.path.join(HERE, "bugs", "BUG-01.png"), width=Mm(240))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()
        para(doc, CAP2, size=10)
        doc.add_picture(os.path.join(HERE, "bugs", "BUG-02.png"), width=Mm(240))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif num == "11":
        for a in ARTIFACTS:
            add_runs(doc.add_paragraph(style="List Bullet"), a, size=Pt(10))

# Bilder zentrieren (add_picture hängt an eigenen Absatz)
for p in doc.paragraphs:
    if p.runs and any(r._r.findall(qn("w:drawing")) for r in p.runs):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

enable_update_fields(doc)
out = os.path.join(HERE, "Aufgabe4_Pruefbericht.docx")
doc.save(out)
print("DOCX:", os.path.relpath(out, HERE), "(%.0f KB)" % (os.path.getsize(out) / 1024))
print("Fertig: professioneller Prüfbericht mit Tabellen + eingebetteten Screenshots.")
